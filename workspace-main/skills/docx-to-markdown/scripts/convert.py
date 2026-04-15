#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
DOCX to Markdown 转换器
支持单文件和批量转换
"""

import argparse
import io
import os
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt

# 修复 Windows 控制台编码问题
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8')


def convert_paragraph(para):
    """将段落转换为 Markdown"""
    text = para.text.strip()
    if not text:
        return ""
    
    # 获取段落样式
    style_name = para.style.name if para.style else ""
    
    # 处理标题
    if "Heading 1" in style_name or style_name.startswith("标题 1"):
        return f"# {text}"
    elif "Heading 2" in style_name or style_name.startswith("标题 2"):
        return f"## {text}"
    elif "Heading 3" in style_name or style_name.startswith("标题 3"):
        return f"### {text}"
    elif "Heading 4" in style_name or style_name.startswith("标题 4"):
        return f"#### {text}"
    elif "Heading 5" in style_name or style_name.startswith("标题 5"):
        return f"##### {text}"
    elif "Heading 6" in style_name or style_name.startswith("标题 6"):
        return f"###### {text}"
    
    # 处理列表
    # 无序列表
    if text.startswith(("•", "·", "-", "○", "●")):
        return f"- {text[1:].strip()}"
    
    # 有序列表（简单匹配）
    if len(text) > 2 and text[0].isdigit():
        if text[1] in ".、)）":
            return f"1. {text[2:].strip()}"
    
    # 处理粗体和斜体（通过 runs 检测）
    # 只有当整个段落都是粗体/斜体时才添加格式标记
    # 避免单个字符被分割加粗的问题
    all_bold = all(run.bold for run in para.runs if run.text.strip())
    all_italic = all(run.italic for run in para.runs if run.text.strip())
    
    if all_bold and all_italic:
        return f"***{text}***"
    elif all_bold:
        return f"**{text}**"
    elif all_italic:
        return f"*{text}*"
    
    return text


def convert_table(table):
    """将表格转换为 Markdown"""
    if not table.rows:
        return ""
    
    lines = []
    rows = list(table.rows)
    
    # 表头
    header_cells = [cell.text.strip() for cell in rows[0].cells]
    lines.append("| " + " | ".join(header_cells) + " |")
    
    # 分隔线
    lines.append("|" + "|".join([" --- " for _ in header_cells]) + "|")
    
    # 数据行
    for row in rows[1:]:
        cells = [cell.text.strip() for cell in row.cells]
        lines.append("| " + " | ".join(cells) + " |")
    
    return "\n".join(lines)


def convert_document(docx_path):
    """转换整个文档"""
    try:
        doc = Document(docx_path)
    except Exception as e:
        raise Exception(f"无法打开文件 {docx_path}: {e}")
    
    md_lines = []
    
    # 处理段落和表格
    for element in doc.element.body:
        tag = element.tag.split("}")[-1]  # 获取标签名
        
        if tag == "p":  # 段落
            # 找到对应的段落对象
            for para in doc.paragraphs:
                if para._element == element:
                    md_line = convert_paragraph(para)
                    if md_line:
                        md_lines.append(md_line)
                    break
        
        elif tag == "tbl":  # 表格
            for table in doc.tables:
                if table._element == element:
                    md_table = convert_table(table)
                    if md_table:
                        md_lines.append("")
                        md_lines.append(md_table)
                        md_lines.append("")
                    break
    
    # 简单遍历方式（如果上述方法有问题）
    if not md_lines:
        # 先处理所有段落
        for para in doc.paragraphs:
            md_line = convert_paragraph(para)
            if md_line:
                md_lines.append(md_line)
        
        # 再处理所有表格
        for table in doc.tables:
            md_table = convert_table(table)
            if md_table:
                md_lines.append("")
                md_lines.append(md_table)
                md_lines.append("")
    
    return "\n".join(md_lines)


def convert_file(input_path, output_path=None):
    """转换单个文件"""
    input_path = Path(input_path)
    
    if not input_path.exists():
        print(f"❌ 错误: 文件不存在: {input_path}")
        return False
    
    if input_path.suffix.lower() != ".docx":
        print(f"❌ 错误: 不是 DOCX 文件: {input_path}")
        return False
    
    # 确定输出路径
    if output_path is None:
        output_path = input_path.with_suffix(".md")
    else:
        output_path = Path(output_path)
    
    print(f"📝 转换: {input_path}")
    print(f"📄 输出: {output_path}")
    
    try:
        md_content = convert_document(input_path)
        
        # 确保输出目录存在
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        # 写入文件
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(md_content)
        
        print(f"✅ 完成: {output_path}")
        return True
        
    except Exception as e:
        print(f"❌ 转换失败: {e}")
        return False


def convert_directory(directory, recursive=False, output_dir=None):
    """批量转换目录"""
    directory = Path(directory)
    
    if not directory.exists():
        print(f"❌ 错误: 目录不存在: {directory}")
        return 0
    
    if not directory.is_dir():
        print(f"❌ 错误: 不是目录: {directory}")
        return 0
    
    # 查找所有 docx 文件
    pattern = "**/*.docx" if recursive else "*.docx"
    docx_files = list(directory.glob(pattern))
    
    if not docx_files:
        print(f"⚠️ 未找到 DOCX 文件: {directory}")
        return 0
    
    print(f"📁 目录: {directory}")
    print(f"🔍 找到 {len(docx_files)} 个 DOCX 文件")
    if recursive:
        print("🔄 递归模式: 包含子目录")
    print()
    
    success_count = 0
    
    for docx_file in docx_files:
        # 确定输出路径
        if output_dir:
            # 保持相对目录结构
            relative_path = docx_file.relative_to(directory)
            output_path = Path(output_dir) / relative_path.with_suffix(".md")
        else:
            output_path = docx_file.with_suffix(".md")
        
        if convert_file(docx_file, output_path):
            success_count += 1
        print()
    
    print(f"📊 转换完成: {success_count}/{len(docx_files)} 成功")
    return success_count


def main():
    parser = argparse.ArgumentParser(
        description="将 Word 文档（DOCX）转换为 Markdown 格式"
    )
    parser.add_argument(
        "input",
        help="输入文件或目录路径"
    )
    parser.add_argument(
        "-o", "--output",
        help="输出文件路径（单文件模式）"
    )
    parser.add_argument(
        "-O", "--output-dir",
        help="输出目录（批量模式）"
    )
    parser.add_argument(
        "-b", "--batch",
        action="store_true",
        help="批量模式：转换目录下所有 DOCX 文件"
    )
    parser.add_argument(
        "-r", "--recursive",
        action="store_true",
        help="递归处理子目录"
    )
    
    args = parser.parse_args()
    
    input_path = Path(args.input)
    
    if args.batch or input_path.is_dir():
        # 批量模式
        convert_directory(
            input_path,
            recursive=args.recursive,
            output_dir=args.output_dir
        )
    else:
        # 单文件模式
        convert_file(args.input, args.output)


if __name__ == "__main__":
    main()
